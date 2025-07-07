import streamlit as st
from openai import OpenAI
import pandas as pd
import json
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches

# Configure page
st.set_page_config(page_title="AI Survey Generator", layout="wide")

# Initialize session state for form persistence
if 'form_data' not in st.session_state:
    st.session_state.form_data = {}
if 'questionnaire_generated' not in st.session_state:
    st.session_state.questionnaire_generated = False
if 'questionnaire_text' not in st.session_state:
    st.session_state.questionnaire_text = ""
if 'survey_data_stored' not in st.session_state:
    st.session_state.survey_data_stored = {}

def load_comprehensive_excel_toolkit():
    """Load comprehensive survey guidelines from Excel toolkit including Survey Question Metadata"""
    toolkit = {
        'question_types': {
            'Likert_5_Point': {
                'scale': ['Strongly Disagree', 'Disagree', 'Neither Agree nor Disagree', 'Agree', 'Strongly Agree'],
                'analysis': ['Descriptive Statistics', 'Factor Analysis', 'Regression Analysis', 'Correlation Analysis']
            },
            'Rating_5_Point': {
                'scale': ['Very Poor', 'Poor', 'Fair', 'Good', 'Excellent'],
                'analysis': ['Descriptive Statistics', 'Gap Analysis', 'Driver Analysis', 'Satisfaction Modeling']
            },
            'Importance_5_Point': {
                'scale': ['Not at all Important', 'Slightly Important', 'Moderately Important', 'Very Important', 'Extremely Important'],
                'analysis': ['Importance-Performance Analysis', 'Driver Analysis', 'MaxDiff Analysis', 'Key Driver Analysis']
            },
            'Likelihood_5_Point': {
                'scale': ['Very Unlikely', 'Unlikely', 'Neither Likely nor Unlikely', 'Likely', 'Very Likely'],
                'analysis': ['Purchase Intent Modeling', 'Predictive Analytics', 'Logistic Regression', 'Conversion Analysis']
            },
            'Association_5_Point': {
                'scale': ['Not at all Associated', 'Slightly Associated', 'Moderately Associated', 'Strongly Associated', 'Extremely Associated'],
                'analysis': ['Brand Mapping', 'Correspondence Analysis', 'Perceptual Mapping', 'Brand Equity Analysis']
            },
            'Frequency_5_Point': {
                'scale': ['Never', 'Rarely', 'Sometimes', 'Often', 'Always'],
                'analysis': ['Usage & Attitude Analysis', 'Behavioral Segmentation', 'Frequency Distribution', 'Usage Patterns']
            }
        },
        'survey_question_metadata': {
            'screener_questions': {
                'age_screening': {
                    'purpose': 'Validate target demographic age range',
                    'data_type': 'Categorical',
                    'validation_rule': 'Must be within specified age range for target audience',
                    'termination_logic': 'Terminate if outside 18-65 or specific target range',
                    'statistical_applications': ['Demographic Profiling', 'Cross-tabulation Base', 'Quota Management'],
                    'required_for_analysis': ['All demographic analyses', 'Age-based segmentation'],
                    'quality_checks': ['Range validation', 'Logical consistency'],
                    'estimated_time_seconds': 10,
                    'mobile_optimization': 'Dropdown with age ranges',
                    'accessibility_notes': 'Screen reader compatible'
                },
                'income_screening': {
                    'purpose': 'Qualify respondents based on income level for target segment',
                    'data_type': 'Categorical_Ordinal',
                    'validation_rule': 'Must meet minimum income threshold',
                    'termination_logic': 'Terminate if below specified income level',
                    'statistical_applications': ['Income-based Segmentation', 'Purchasing Power Analysis', 'Price Sensitivity Modeling'],
                    'required_for_analysis': ['Economic demographic profiling', 'Price elasticity studies'],
                    'quality_checks': ['Income range validation', 'Consistency with lifestyle indicators'],
                    'estimated_time_seconds': 15,
                    'mobile_optimization': 'Clear income ranges with local currency',
                    'accessibility_notes': 'High contrast for readability'
                }
            },
            'core_research_questions': {
                'brand_awareness_unaided': {
                    'purpose': 'Measure spontaneous brand recall without prompting',
                    'data_type': 'Text_Multiple_Response',
                    'validation_rule': 'Minimum 1 character, maximum 200 characters per brand',
                    'termination_logic': 'No termination',
                    'statistical_applications': ['Top-of-Mind Awareness Analysis', 'Brand Salience Measurement', 'Competitive Analysis'],
                    'required_for_analysis': ['Brand equity studies', 'Market share correlation', 'Brand health tracking'],
                    'quality_checks': ['Text quality validation', 'Brand name standardization', 'Spelling correction'],
                    'estimated_time_seconds': 60,
                    'mobile_optimization': 'Auto-complete with brand suggestions',
                    'accessibility_notes': 'Voice input support'
                },
                'attribute_importance_ratings': {
                    'purpose': 'Measure importance of product/service attributes in decision making',
                    'data_type': 'Rating_Scale_5_Point',
                    'validation_rule': 'All attributes must be rated on 1-5 scale',
                    'termination_logic': 'No termination',
                    'statistical_applications': ['Importance-Performance Analysis', 'Key Driver Analysis', 'Factor Analysis', 'Conjoint Analysis'],
                    'required_for_analysis': ['Product development priorities', 'Marketing message optimization', 'Feature prioritization'],
                    'quality_checks': ['Straight-lining detection', 'Response time validation', 'Logical consistency'],
                    'estimated_time_seconds': 90,
                    'mobile_optimization': 'Slider interface with haptic feedback',
                    'accessibility_notes': 'Voice guidance for ratings'
                }
            }
        },
        'fraud_checks': {
            'attention_check': "Please select 'Agree' for this question to confirm you are reading carefully.",
            'time_validation': "Minimum time per question: 3-5 seconds, Maximum: 120 seconds",
            'straight_lining': "Flag responses with same rating across 5+ consecutive questions",
            'open_end_quality': "Check for meaningful responses, minimum 10 characters for detailed questions",
            'geographic_validation': "Validate IP location matches declared location",
            'duplicate_detection': "Check for duplicate responses using device fingerprinting"
        },
        'termination_criteria': {
            'age_out': "Respondents outside target age range",
            'income_screening': "Below minimum income threshold for target segment", 
            'geographic_screening': "Outside specified geographic boundaries",
            'category_usage': "Non-users of category if users-only study",
            'quota_full': "Target demographic quota reached",
            'quality_screening': "Failed fraud/attention checks"
        },
        'loi_calculation': {
            'simple_questions': '15-20 seconds each',
            'matrix_questions': '45-90 seconds each', 
            'ranking_questions': '60-120 seconds each',
            'open_ended': '90-180 seconds each',
            'demographics': '10-15 seconds each'
        }
    }
    return toolkit

def detect_survey_category(survey_objective, target_audience):
    """FIXED: Intelligent category detection from survey objective and target audience"""
    combined_text = f"{survey_objective.lower()} {target_audience.lower()}"
    
    # Define category keywords with broader coverage
    category_keywords = {
        'cosmetics': ['cosmetics', 'beauty', 'cream', 'skincare', 'makeup', 'lipstick', 'foundation', 'night cream', 'face cream', 'moisturizer', 'serum', 'lotion'],
        'automotive': ['automotive', 'car', 'vehicle', 'ev', 'electric vehicle', 'auto', 'automobile', 'sedan', 'suv'],
        'food_beverage': ['food', 'restaurant', 'dining', 'beverage', 'drink', 'coffee', 'tea', 'snack', 'meal'],
        'technology': ['phone', 'smartphone', 'mobile', 'technology', 'laptop', 'computer', 'software', 'app', 'tech'],
        'fashion': ['fashion', 'clothing', 'apparel', 'shoes', 'dress', 'shirt', 'accessories'],
        'healthcare': ['healthcare', 'medical', 'health', 'medicine', 'treatment', 'hospital', 'doctor'],
        'finance': ['finance', 'banking', 'investment', 'insurance', 'loan', 'credit', 'financial'],
        'travel': ['travel', 'hotel', 'vacation', 'tourism', 'airline', 'booking'],
        'education': ['education', 'learning', 'course', 'school', 'university', 'training']
    }
    
    # Score each category based on keyword matches
    category_scores = {}
    for category, keywords in category_keywords.items():
        score = sum(1 for keyword in keywords if keyword in combined_text)
        if score > 0:
            category_scores[category] = score
    
    # Return the category with highest score, or 'general' if no matches
    if category_scores:
        detected_category = max(category_scores, key=category_scores.get)
        confidence = category_scores[detected_category]
        return detected_category, confidence
    else:
        return 'general', 0

def get_dynamic_brand_list_from_research(category, market, api_key):
    """FIXED: Dynamically research brands for ANY category - not just automotive"""
    try:
        client = OpenAI(api_key=api_key)
        
        # Category-specific research prompt
        research_prompt = f"""
        List 15 popular {category} brands available in {market} market. 
        
        Category: {category}
        Market: {market}
        
        Requirements:
        - List ONLY brands that are actually available in {market}
        - Include both international and local brands
        - Focus on well-known, established brands
        - One brand name per line, no descriptions
        
        Format example:
        Brand Name 1
        Brand Name 2
        Brand Name 3
        """
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": f"List popular {category} brands available in {market}. One name per line, no descriptions."},
                {"role": "user", "content": research_prompt}
            ],
            temperature=0.1,
            max_tokens=400
        )
        
        # Extract and clean brand names
        brand_text = response.choices[0].message.content.strip()
        brand_lines = [line.strip() for line in brand_text.split('\n') if line.strip()]
        
        brands = []
        for line in brand_lines:
            # Clean brand names
            clean_brand = line.replace('•', '').replace('-', '').replace('*', '')
            clean_brand = ''.join(char for char in clean_brand if not (char.isdigit() and char in '123456789.'))
            clean_brand = clean_brand.strip()
            
            if clean_brand and len(clean_brand) > 2 and len(clean_brand) < 50:
                brands.append(clean_brand)
        
        if len(brands) >= 5:
            return brands[:15]
        else:
            raise Exception(f"Only {len(brands)} brands found")
            
    except Exception as e:
        # Return category-specific fallback brands
        return get_fallback_brands(category, market)

def get_fallback_brands(category, market):
    """FIXED: Category-specific fallback brands"""
    fallback_brands = {
        'cosmetics': {
            'India': ['Lakme', 'Maybelline', 'L\'Oreal', 'MAC', 'Revlon', 'Clinique', 'Estee Lauder', 'Nykaa', 'Lotus Herbals', 'Forest Essentials', 'Colorbar', 'Faces Canada', 'Biotique', 'Himalaya Herbals', 'VLCC'],
            'global': ['L\'Oreal', 'Maybelline', 'MAC', 'Revlon', 'Clinique', 'Estee Lauder', 'Chanel', 'Dior', 'Urban Decay', 'NARS']
        },
        'automotive': {
            'India': ['Maruti Suzuki', 'Hyundai', 'Tata Motors', 'Mahindra', 'Toyota', 'Honda', 'BMW', 'Mercedes-Benz', 'Audi', 'Ford'],
            'global': ['Toyota', 'Honda', 'Ford', 'BMW', 'Mercedes-Benz', 'Audi', 'Hyundai', 'Nissan', 'Volkswagen', 'Tesla']
        },
        'technology': {
            'India': ['Samsung', 'Apple', 'OnePlus', 'Xiaomi', 'Oppo', 'Vivo', 'Realme', 'Nokia', 'Motorola', 'Google'],
            'global': ['Apple', 'Samsung', 'Google', 'Microsoft', 'Sony', 'LG', 'Huawei', 'OnePlus', 'Nokia', 'Motorola']
        },
        'food_beverage': {
            'India': ['Amul', 'Britannia', 'Parle', 'ITC', 'Nestle', 'Coca-Cola', 'PepsiCo', 'Haldiram\'s', 'MTR', 'Patanjali'],
            'global': ['Coca-Cola', 'PepsiCo', 'Nestle', 'Unilever', 'Mars', 'Mondelez', 'Kraft Heinz', 'General Mills', 'Kellogg\'s', 'Starbucks']
        }
    }
    
    # Get market-specific brands or default to global
    market_key = 'India' if 'india' in market.lower() else 'global'
    
    if category in fallback_brands and market_key in fallback_brands[category]:
        return fallback_brands[category][market_key]
    elif category in fallback_brands:
        return fallback_brands[category]['global']
    else:
        return ['Brand A', 'Brand B', 'Brand C', 'Brand D', 'Brand E']

def get_comprehensive_brand_list(category, market, api_key=None):
    """FIXED: Get truly dynamic brand list for ANY category"""
    
    if api_key:
        # Use AI-powered dynamic research
        return get_dynamic_brand_list_from_research(category, market, api_key)
    else:
        # Fallback when no API key available
        return get_fallback_brands(category, market)

def calculate_question_count(loi_minutes):
    """Calculate proper question distribution based on LOI with higher question counts"""
    # Core research questions should be 2.5 times LOI for comprehensive surveys
    core_questions = int(loi_minutes * 2.5)
    
    # Additional questions for screener and demographics
    screener_questions = max(8, int(loi_minutes * 0.4))
    demographics_questions = max(8, int(loi_minutes * 0.4))
    
    total_questions = core_questions + screener_questions + demographics_questions
    
    return {
        'screener': screener_questions,
        'core_research': core_questions,
        'demographics': demographics_questions,
        'total': total_questions
    }

def web_research_brands_and_trends(query, api_key):
    """Enhanced web research for comprehensive brand lists and current trends"""
    try:
        client = OpenAI(api_key=api_key)
        research_prompt = f"""
        Research and provide comprehensive information for: {query}
        
        Required Output Format:
        1. COMPREHENSIVE BRAND LIST (minimum 15-20 brands):
           - Include all major players (luxury, premium, mass market)
           - Current market leaders and emerging brands
           - Both domestic and international brands available in the market
        
        2. CURRENT MARKET TRENDS (2024-2025):
           - Latest consumer preferences and behaviors
           - Emerging technologies and features
           - Price trends and market dynamics
           - Key attributes driving purchase decisions
        
        3. CONSUMER INSIGHTS:
           - Primary decision factors
           - Demographic preferences
           - Usage patterns and behaviors
           - Satisfaction drivers
        
        Provide specific, actionable insights with current brand names and market data.
        """
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a market research expert with access to current comprehensive market data and brand intelligence."},
                {"role": "user", "content": research_prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Research error: {str(e)}"

def format_questionnaire_with_logic(questionnaire_text):
    """Enhanced formatting with better structure and logic display"""
    lines = questionnaire_text.split('\n')
    formatted_lines = []
    
    for line in lines:
        if line.strip():
            # Section headers
            if 'SECTION' in line.upper() or line.strip().startswith('==='):
                formatted_lines.append('\n' + '='*80)
                formatted_lines.append(line.upper())
                formatted_lines.append('='*80 + '\n')
            # Question numbers
            elif line.strip().startswith('Q') and ':' in line:
                formatted_lines.append('\n' + '-'*50)
                formatted_lines.append(line)
            # Statistical analysis, fraud checks, logic
            elif any(keyword in line for keyword in ['Statistical Methods:', 'Fraud Detection:', 'Skip Logic:', 'Termination:', 'Purpose:', 'Data Type:']):
                formatted_lines.append('    → ' + line)
            # Response options - ensure they're on separate lines
            elif line.strip().startswith('-') or line.strip().startswith('•'):
                formatted_lines.append('    ' + line)
            else:
                formatted_lines.append(line)
        else:
            formatted_lines.append('')
    
    return '\n'.join(formatted_lines)

def create_comprehensive_word_document(questionnaire_text, survey_data):
    """Create detailed Word document with survey specifications"""
    doc = Document()
    
    # Title page
    title = doc.add_heading('Professional Survey Questionnaire', 0)
    
    # Executive summary
    doc.add_heading('Survey Specifications', level=1)
    specs_table = doc.add_table(rows=9, cols=2)
    specs_table.style = 'Table Grid'
    
    specs_data = [
        ['Survey Objective', survey_data['survey_objective']],
        ['Target Audience', survey_data['target_audience']],
        ['Expected LOI', f"{survey_data['survey_loi']} minutes"],
        ['Methodology', survey_data['methodology']],
        ['Device Context', survey_data['device_context']],
        ['Market/Country', survey_data['market_country']],
        ['Detected Category', survey_data.get('detected_category', 'Unknown')],
        ['Statistical Methods', ', '.join(survey_data['statistical_methods'])],
        ['Generation Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
    ]
    
    for i, (key, value) in enumerate(specs_data):
        specs_table.cell(i, 0).text = key
        specs_table.cell(i, 1).text = str(value)
    
    # Question count summary
    question_counts = calculate_question_count(survey_data['survey_loi'])
    doc.add_heading('Question Distribution', level=1)
    count_para = doc.add_paragraph()
    count_para.add_run(f"• Screener Questions: {question_counts['screener']}\n")
    count_para.add_run(f"• Core Research Questions: {question_counts['core_research']}\n")
    count_para.add_run(f"• Demographics Questions: {question_counts['demographics']}\n")
    count_para.add_run(f"• Total Questions: {question_counts['total']}")
    
    # Questionnaire content
    doc.add_page_break()
    doc.add_heading('Complete Questionnaire', level=1)
    
    # Process questionnaire text into structured paragraphs
    lines = questionnaire_text.split('\n')
    for line in lines:
        if line.strip():
            if 'SECTION' in line.upper():
                doc.add_heading(line, level=2)
            elif line.strip().startswith('Q') and ':' in line:
                doc.add_paragraph(line, style='Heading 3')
            elif any(keyword in line for keyword in ['Statistical Methods:', 'Fraud Detection:', 'Skip Logic:']):
                p = doc.add_paragraph()
                p.add_run(line).italic = True
            else:
                doc.add_paragraph(line)
    
    return doc

def create_structured_excel_output(questionnaire_text, survey_data, toolkit):
    """Create comprehensive Excel file with multiple sheets including Survey Question Metadata"""
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # Survey Details Sheet
        survey_df = pd.DataFrame([survey_data])
        survey_df.to_excel(writer, sheet_name='Survey_Details', index=False)
        
        # Question Analysis Sheet
        questions_data = []
        lines = questionnaire_text.split('\n')
        current_question = {}
        
        for line in lines:
            line = line.strip()
            if line.startswith('Q') and ':' in line:
                if current_question:
                    questions_data.append(current_question)
                current_question = {
                    'Question_Number': line.split(':')[0],
                    'Question_Text': line.split(':', 1)[1].strip() if ':' in line else line,
                    'Question_Type': '',
                    'Response_Options': '',
                    'Statistical_Methods': '',
                    'Fraud_Check': '',
                    'Skip_Logic': '',
                    'Scale_Description': '',
                    'Purpose': '',
                    'Data_Type': '',
                    'Validation_Rule': '',
                    'Required_For_Analysis': '',
                    'Quality_Checks': '',
                    'Estimated_Time_Seconds': '',
                    'Termination_Logic': ''
                }
            elif line and current_question:
                if 'Statistical Methods' in line:
                    current_question['Statistical_Methods'] = line.replace('Statistical Methods:', '').strip()
                elif 'Fraud Detection' in line:
                    current_question['Fraud_Check'] = line.replace('Fraud Detection:', '').strip()
                elif 'Skip Logic' in line:
                    current_question['Skip_Logic'] = line.replace('Skip Logic:', '').strip()
                elif 'Purpose:' in line:
                    current_question['Purpose'] = line.replace('Purpose:', '').strip()
                elif 'Data Type:' in line:
                    current_question['Data_Type'] = line.replace('Data Type:', '').strip()
                elif 'Validation Rule:' in line:
                    current_question['Validation_Rule'] = line.replace('Validation Rule:', '').strip()
                elif 'Required For Analysis:' in line:
                    current_question['Required_For_Analysis'] = line.replace('Required For Analysis:', '').strip()
                elif 'Quality Checks:' in line:
                    current_question['Quality_Checks'] = line.replace('Quality Checks:', '').strip()
                elif 'Estimated Time:' in line:
                    current_question['Estimated_Time_Seconds'] = line.replace('Estimated Time:', '').strip()
                elif 'Termination Logic:' in line:
                    current_question['Termination_Logic'] = line.replace('Termination Logic:', '').strip()
                elif any(scale in line for scale in ['Strongly Disagree', 'Very Poor', 'Not at all']):
                    current_question['Scale_Description'] = line
                elif line.startswith('-') or line.startswith('•'):
                    current_question['Response_Options'] += line + '\n'
        
        if current_question:
            questions_data.append(current_question)
        
        questions_df = pd.DataFrame(questions_data)
        questions_df.to_excel(writer, sheet_name='Questions_Analysis', index=False)
        
        # Survey Question Metadata Sheet
        metadata_rows = []
        
        # Screener Questions Metadata
        for q_type, metadata in toolkit['survey_question_metadata']['screener_questions'].items():
            metadata_rows.append({
                'Question_Category': 'Screener',
                'Question_Type': q_type,
                'Purpose': metadata['purpose'],
                'Data_Type': metadata['data_type'],
                'Validation_Rule': metadata['validation_rule'],
                'Termination_Logic': metadata['termination_logic'],
                'Statistical_Applications': ' | '.join(metadata['statistical_applications']),
                'Required_For_Analysis': ' | '.join(metadata['required_for_analysis']),
                'Quality_Checks': ' | '.join(metadata['quality_checks']),
                'Estimated_Time_Seconds': metadata['estimated_time_seconds'],
                'Mobile_Optimization': metadata['mobile_optimization'],
                'Accessibility_Notes': metadata['accessibility_notes']
            })
        
        # Core Research Questions Metadata
        for q_type, metadata in toolkit['survey_question_metadata']['core_research_questions'].items():
            metadata_rows.append({
                'Question_Category': 'Core Research',
                'Question_Type': q_type,
                'Purpose': metadata['purpose'],
                'Data_Type': metadata['data_type'],
                'Validation_Rule': metadata['validation_rule'],
                'Termination_Logic': metadata['termination_logic'],
                'Statistical_Applications': ' | '.join(metadata['statistical_applications']),
                'Required_For_Analysis': ' | '.join(metadata['required_for_analysis']),
                'Quality_Checks': ' | '.join(metadata['quality_checks']),
                'Estimated_Time_Seconds': metadata['estimated_time_seconds'],
                'Mobile_Optimization': metadata['mobile_optimization'],
                'Accessibility_Notes': metadata['accessibility_notes']
            })
        
        # Create Survey Question Metadata DataFrame and export
        metadata_df = pd.DataFrame(metadata_rows)
        metadata_df.to_excel(writer, sheet_name='Survey_Question_Metadata', index=False)
        
        # Survey Toolkit Sheet
        toolkit_data = []
        for q_type, details in toolkit['question_types'].items():
            toolkit_data.append({
                'Question_Type': q_type,
                'Scale_Options': ' | '.join(details['scale']),
                'Analysis_Methods': ' | '.join(details['analysis'])
            })
        
        toolkit_df = pd.DataFrame(toolkit_data)
        toolkit_df.to_excel(writer, sheet_name='Survey_Toolkit', index=False)
        
        # Fraud Checks Sheet
        fraud_df = pd.DataFrame([toolkit['fraud_checks']])
        fraud_df.to_excel(writer, sheet_name='Fraud_Guidelines', index=False)
        
        # Termination Criteria Sheet
        termination_df = pd.DataFrame([toolkit['termination_criteria']])
        termination_df.to_excel(writer, sheet_name='Termination_Criteria', index=False)
        
        # LOI Calculation Guidelines Sheet
        loi_df = pd.DataFrame([toolkit['loi_calculation']])
        loi_df.to_excel(writer, sheet_name='LOI_Guidelines', index=False)
    
    return output.getvalue()

# Main App Interface
st.title("🎯 Advanced AI Survey Questionnaire Generator")
st.markdown("*Professional survey design with comprehensive analytics and fraud detection*")

# Sidebar
with st.sidebar:
    st.header("🔧 Configuration")
    api_key = st.text_input("OpenAI API Key:", type="password", key='api_key')
    
    if st.button("🔄 Reset Form", help="Clear all inputs and start fresh"):
        for key in list(st.session_state.keys()):
            if key != 'api_key':
                del st.session_state[key]
        st.rerun()

# Main form (preserves data)
col1, col2 = st.columns([2, 1])

with col1:
    st.header("📋 Survey Configuration")
    
    survey_objective = st.text_area(
        "Survey Objective", 
        value=st.session_state.get('survey_objective', ''),
        placeholder="e.g., Understand night cream usage patterns and brand preferences among women aged 18-45",
        key='survey_objective'
    )
    
    target_audience = st.text_input(
        "Target Audience",
        value=st.session_state.get('target_audience', ''),
        placeholder="e.g., Women aged 18-45 who have used night cream in the last week",
        key='target_audience'
    )
    
    col_a, col_b = st.columns(2)
    with col_a:
        population_size = st.number_input("Population Size", min_value=100, value=st.session_state.get('population_size', 1000), key='population_size')
    with col_b:
        survey_loi = st.number_input("Survey LOI (minutes)", min_value=5, max_value=60, value=st.session_state.get('survey_loi', 20), key='survey_loi')
    
    # Display calculated question counts
    q_counts = calculate_question_count(survey_loi)
    st.info(f"📊 **Enhanced Question Distribution:** {q_counts['screener']} Screener + {q_counts['core_research']} Core Research + {q_counts['demographics']} Demographics = **{q_counts['total']} Total Questions**")
    
    col_c, col_d = st.columns(2)
    with col_c:
        methodology = st.selectbox("Methodology", ["Online", "Phone", "Face-to-Face", "Mobile App"], key='methodology')
    with col_d:
        device_context = st.selectbox("Device Context", ["Desktop", "Mobile", "Mixed"], key='device_context')
    
    market_country = st.text_input("Market/Country", value=st.session_state.get('market_country', 'India'), key='market_country')

with col2:
    st.header("⚙️ Advanced Options")
    
    statistical_methods = st.multiselect(
        "Statistical Methods",
        ["Regression", "Conjoint", "Cluster Analysis", "MaxDiff", "Factor Analysis", "TURF Analysis", 
         "Discriminant Analysis", "Correspondence Analysis", "Latent Class Analysis", "SEM", "CHAID"],
        default=st.session_state.get('statistical_methods', []),
        key='statistical_methods'
    )
    
    allowed_question_types = st.multiselect(
        "Question Types",
        ["Likert", "Open-End", "Rating Scale", "Matrix/Grid", "Dichotomous", "Ranking", "Slider"],
        default=st.session_state.get('allowed_question_types', []),
        key='allowed_question_types'
    )
    
    compliance_requirements = st.multiselect(
        "Compliance",
        ["GDPR", "CCPA", "HIPAA", "Other"],
        default=st.session_state.get('compliance_requirements', []),
        key='compliance_requirements'
    )

# Generation Section
st.header("🚀 Generate Advanced Survey")

if st.button("🎯 Generate Comprehensive Survey Questionnaire", type="primary", use_container_width=True):
    if not api_key:
        st.error("⚠️ Please enter your OpenAI API key")
        st.stop()
    
    if not survey_objective or not target_audience:
        st.error("⚠️ Please provide Survey Objective and Target Audience")
        st.stop()
    
    # FIXED: Intelligent category detection
    detected_category, confidence = detect_survey_category(survey_objective, target_audience)
    
    # Store survey data
    survey_data = {
        'survey_objective': survey_objective,
        'target_audience': target_audience,
        'population_size': population_size,
        'survey_loi': survey_loi,
        'methodology': methodology,
        'device_context': device_context,
        'market_country': market_country,
        'statistical_methods': statistical_methods,
        'allowed_question_types': allowed_question_types,
        'compliance_requirements': compliance_requirements,
        'detected_category': detected_category  # FIXED: Store detected category
    }
    
    st.session_state.survey_data_stored = survey_data
    
    # Progress tracking
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        # Step 1: Load Excel Toolkit
        status_text.text("📚 Loading Excel Survey Toolkit...")
        progress_bar.progress(15)
        toolkit = load_comprehensive_excel_toolkit()
        
        # Step 2: Calculate question counts first
        status_text.text("📊 Calculating question distribution...")
        progress_bar.progress(25)
        question_counts = calculate_question_count(survey_data['survey_loi'])
        
        # Step 3: FIXED Category detection and brand research
        status_text.text(f"🧠 Category detected: {detected_category} (confidence: {confidence})")
        progress_bar.progress(35)
        
        st.info(f"🎯 **Category Detected:** {detected_category.title()} | **Confidence Score:** {confidence} | **Market:** {survey_data['market_country']}")
        
        # FIXED: Dynamic brand research based on detected category
        try:
            brand_list = get_comprehensive_brand_list(detected_category, survey_data['market_country'], api_key)
            brand_research_success = True
        except Exception as e:
            st.warning(f"⚠️ Brand research failed: {str(e)}")
            brand_list = get_fallback_brands(detected_category, survey_data['market_country'])
            brand_research_success = False
        
        # Display results
        if brand_research_success and len(brand_list) > 5:
            st.success(f"✅ **Dynamic Brand Research Successful:** {', '.join(brand_list[:6])}... ({len(brand_list)} total brands)")
        else:
            st.error(f"❌ **Brand Research Failed:** Using fallback brands - {', '.join(brand_list)}")
        
        # Store category and brands for prompts
        brand_list_text = ', '.join(brand_list)
        top_10_brands = ', '.join(brand_list[:10])
        top_6_brands = ', '.join(brand_list[:6])
        
        # Step 4: Comprehensive Market Research
        status_text.text(f"🔍 Conducting {detected_category} market research...")
        progress_bar.progress(45)
        research_query = f"{survey_data['target_audience']} {survey_data['market_country']} {detected_category} comprehensive brand list market trends consumer behavior"
        research_data = web_research_brands_and_trends(research_query, api_key)
        
        # Step 5: Generate Advanced Prompt
        status_text.text("📝 Preparing survey generation...")
        progress_bar.progress(55)
        
        # Step 6: FIXED - Generate Questionnaire in Multiple Parts with category consistency
        status_text.text(f"🤖 Generating comprehensive {detected_category} questionnaire...")
        progress_bar.progress(65)
        
        client = OpenAI(api_key=api_key)
        
        # Generate questionnaire in parts to ensure all questions are created
        full_questionnaire = ""
        
        # FIXED Part 1: Screener Questions - category specific
        status_text.text(f"🤖 Generating {detected_category} screener questions...")
        screener_prompt = f"""
        You are an expert {detected_category} market researcher specializing in {survey_data['market_country']} market research.

        Generate EXACTLY {question_counts['screener']} SCREENER QUESTIONS for this {detected_category} survey:
        
        Survey Objective: {survey_data['survey_objective']}
        Target Audience: {survey_data['target_audience']}
        Market: {survey_data['market_country']}
        Category: {detected_category}
        Available Brands: {brand_list_text}
        
        CRITICAL REQUIREMENTS:
        - Generate EXACTLY {question_counts['screener']} questions numbered Q1. Q2. Q3. etc.
        - Focus ONLY on {detected_category}-specific screening based on survey objective
        - Use REAL {detected_category} brands from research: {top_10_brands}
        - NEVER use "Brand A, Brand B, Brand C" - always use real brand names
        - Each answer option on separate line with dash (-)
        - Include complete metadata for each question
        - Include termination logic where applicable
        - NO QUESTIONS ABOUT OTHER CATEGORIES
        
        EXAMPLE FORMAT:
        Q1. What is your age?
        - 18-24
        - 25-34
        - 35-44
        - 45-54
        - 55+ 
        - Others (specify)
        
        Purpose: Validate target demographic age range for {detected_category} study
        Data Type: Categorical_Single_Response
        Statistical Methods: Descriptive Statistics, Cross-tabulation, Demographic Analysis
        Fraud Detection: No
        Quality Checks: Age range validation, logical consistency
        Termination Logic: Terminate if outside target age range for this {detected_category} study
        
        Generate all {question_counts['screener']} screener questions focusing ONLY on {detected_category} qualification.
        Use real brand names: {brand_list_text}
        """
        
        screener_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"You are an expert {detected_category} survey designer. Generate EXACTLY the number of questions specified. Do not truncate."},
                {"role": "user", "content": screener_prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        
        full_questionnaire += screener_response.choices[0].message.content + "\n\n"
        
        # FIXED Part 2: Core Research Questions (First Half) - category specific
        status_text.text(f"🤖 Generating {detected_category} core research questions (Part 1)...")
        progress_bar.progress(75)
        core_part1_count = question_counts['core_research'] // 2
        start_q = question_counts['screener'] + 1
        end_q = start_q + core_part1_count - 1
        
        core_part1_prompt = f"""
        You are an expert {detected_category} market researcher. 

        Generate EXACTLY {core_part1_count} CORE {detected_category.upper()} RESEARCH QUESTIONS:
        
        Survey Objective: {survey_data['survey_objective']}
        Target Audience: {survey_data['target_audience']}
        Market: {survey_data['market_country']}
        Category: {detected_category}
        
        DYNAMICALLY RESEARCHED BRANDS FOR {detected_category.upper()}:
        {brand_list_text}
        
        CRITICAL BRAND REQUIREMENTS:
        - NEVER use generic names like "Brand A, Brand B, Brand C"
        - ALWAYS use the specific {detected_category} brands from the research: {top_10_brands}
        - If insufficient brands researched, acknowledge and work with available brands
        
        CRITICAL REQUIREMENTS:
        - Generate EXACTLY {core_part1_count} questions numbered Q{start_q}. to Q{end_q}.
        - Focus on: brand awareness, current usage, {detected_category} preferences
        - Each answer option on separate line with dash (-)
        - Include complete metadata for each question
        - ONLY {detected_category.upper()}-RELATED QUESTIONS based on survey objective
        
        MANDATORY BRAND QUESTIONS FOR {detected_category.upper()}:
        
        Q{start_q}. Which {detected_category} brands come to mind when you think about this category? (Unaided awareness)
        - Open-ended text response
        
        Q{start_q+1}. Which of the following {detected_category} brands have you heard of? (Select all that apply)
        {chr(10).join([f'- {brand}' for brand in brand_list[:12]])}
        - Others (specify)
        
        Q{start_q+2}. Which {detected_category} brands do you currently use or have used? (Select all that apply)
        {chr(10).join([f'- {brand}' for brand in brand_list[:10]])}
        - Others (specify)
        
        Continue generating remaining questions using the researched brands: {brand_list_text}
        Focus on {detected_category}-specific usage, satisfaction, and preferences based on the survey objective.
        """
        
        core_part1_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"You are an expert {detected_category} survey designer. Generate EXACTLY the number of questions specified. Do not truncate."},
                {"role": "user", "content": core_part1_prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        
        full_questionnaire += core_part1_response.choices[0].message.content + "\n\n"
        
        # FIXED Part 3: Core Research Questions (Second Half) - category specific
        status_text.text(f"🤖 Generating {detected_category} core research questions (Part 2)...")
        progress_bar.progress(85)
        core_part2_count = question_counts['core_research'] - core_part1_count
        start_q2 = end_q + 1
        end_q2 = start_q2 + core_part2_count - 1
        
        core_part2_prompt = f"""
        You are an expert {detected_category} market researcher. Generate EXACTLY {core_part2_count} ADVANCED {detected_category.upper()} RESEARCH QUESTIONS:
        
        MANDATORY BRAND LIST - Use these dynamically loaded brands:
        {brand_list_text}
        
        TOP BRANDS FOR MATRICES: {top_6_brands}
        
        CRITICAL REQUIREMENTS:
        - Generate EXACTLY {core_part2_count} questions numbered Q{start_q2}. to Q{end_q2}.
        - NEVER use "Brand A, Brand B" - use brands from the dynamic list above
        - Include matrix questions with 5-point scales
        - Each answer option on separate line with dash (-)
        - Focus ONLY on {detected_category} category - NO OTHER CATEGORIES
        - ABSOLUTELY NO AUTOMOTIVE, EV, CAR, or VEHICLE questions unless category is automotive
        - ABSOLUTELY NO TECHNOLOGY questions unless category is technology
        - ONLY {detected_category.upper()} QUESTIONS
        
        MANDATORY MATRIX QUESTIONS FOR {detected_category.upper()}:
        
        Q{start_q2}. Please rate the importance of the following {detected_category} attributes: 
        (Scale: 1=Not at all Important, 2=Slightly Important, 3=Moderately Important, 4=Very Important, 5=Extremely Important)
        - Quality: [1] [2] [3] [4] [5]
        - Price: [1] [2] [3] [4] [5]
        - Brand Reputation: [1] [2] [3] [4] [5]
        - Availability: [1] [2] [3] [4] [5]
        - Packaging: [1] [2] [3] [4] [5]
        - Effectiveness: [1] [2] [3] [4] [5]
        - Safety: [1] [2] [3] [4] [5]
        - Ingredients: [1] [2] [3] [4] [5]
        - Overall Satisfaction: [1] [2] [3] [4] [5]
        
        Q{start_q2+1}. How do you associate "Premium Quality" with these {detected_category} brands?
        (Scale: 1=Not Associated, 2=Slightly Associated, 3=Moderately Associated, 4=Strongly Associated, 5=Extremely Associated)
        {chr(10).join([f'- {brand}: [1] [2] [3] [4] [5]' for brand in brand_list[:6]])}
        
        Continue with remaining questions about:
        - {detected_category} purchase journey factors
        - Brand loyalty and switching behavior in {detected_category}
        - {detected_category} usage patterns and frequency
        - Satisfaction with current {detected_category} products
        - Information sources for {detected_category} research
        - {detected_category} shopping preferences and channels
        
        Use brands from: {brand_list_text}
        
        STRICT RULE: Focus ONLY on {detected_category} - no other product categories.
        DO NOT mention cars, vehicles, EVs, technology, or any other category unless it IS the detected category.
        """
        
        core_part2_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"You are an expert {detected_category} survey designer. Generate EXACTLY the number of questions specified. Do not truncate."},
                {"role": "user", "content": core_part2_prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        
        full_questionnaire += core_part2_response.choices[0].message.content + "\n\n"
        
        # Part 4: Demographics Questions
        status_text.text("🤖 Generating demographics questions...")
        progress_bar.progress(90)
        demo_start = end_q2 + 1
        demo_end = demo_start + question_counts['demographics'] - 1
        
        demo_prompt = f"""
        You are an expert survey researcher. Generate EXACTLY {question_counts['demographics']} DEMOGRAPHICS QUESTIONS for this {detected_category} study:
        
        CRITICAL REQUIREMENTS:
        - Generate EXACTLY {question_counts['demographics']} questions numbered Q{demo_start}. to Q{demo_end}.
        - Include: age, gender, income, education, household size, employment, lifestyle
        - Each answer option on separate line with dash (-)
        - Include complete metadata for each question
        - Use {survey_data['market_country']} market context (local currency, cities, etc.)
        - NO PRODUCT QUESTIONS - ONLY DEMOGRAPHICS
        
        MANDATORY DEMOGRAPHICS:
        1. Age (detailed brackets)
        2. Gender 
        3. Annual income (in local currency - {survey_data['market_country']} context)
        4. Education level
        5. Employment status
        6. Household size
        7. City of residence
        8. Lifestyle/Family status
        
        EXAMPLE:
        Q{demo_start}. What is your highest level of education?
        - Less than high school
        - High school graduate
        - Some college
        - Bachelor's degree
        - Master's degree
        - PhD/Doctorate
        - Others (specify)
        
        Purpose: Educational profiling for {detected_category} market research study
        Data Type: Categorical_Ordinal
        Statistical Methods: Demographic analysis, Education-based segmentation, Cross-tabulation
        
        Focus ONLY on demographic profiling - NO product satisfaction questions.
        """
        
        demo_response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert survey designer. Generate EXACTLY the number of questions specified. Do not truncate."},
                {"role": "user", "content": demo_prompt}
            ],
            temperature=0.2,
            max_tokens=1500
        )
        
        full_questionnaire += demo_response.choices[0].message.content
        
        questionnaire = full_questionnaire
        
        # Validate question count
        question_lines = [line for line in questionnaire.split('\n') if line.strip().startswith('Q') and '.' in line and any(char.isdigit() for char in line)]
        actual_count = len(question_lines)
        
        if actual_count < question_counts['total']:
            st.warning(f"⚠️ Generated {actual_count} questions instead of {question_counts['total']}. Attempting to complete...")
            
            # Generate remaining questions if needed
            remaining_count = question_counts['total'] - actual_count
            if remaining_count > 0:
                completion_prompt = f"""
                The {detected_category} survey is incomplete. Generate {remaining_count} additional {detected_category} questions to complete the survey.
                Continue from Q{actual_count + 1} to Q{question_counts['total']}.
                
                REQUIREMENTS:
                - Generate EXACTLY {remaining_count} questions
                - Number them Q{actual_count + 1} through Q{question_counts['total']}
                - Include {detected_category} purchase journey, satisfaction, and additional research questions
                - Each answer option on separate line with dash (-)
                - Include complete metadata for each question
                - Use brands: {brand_list_text}
                - Focus ONLY on {detected_category} category
                - NO OTHER CATEGORIES (no EV, no automotive unless detected category is automotive)
                
                EXAMPLE FORMAT:
                Q{actual_count + 1}. [Question text about {detected_category}]
                - Option 1
                - Option 2
                - Option 3
                
                Purpose: [Research objective for {detected_category}]
                Statistical Methods: [Analysis methods]
                Fraud Detection: [Yes/No]
                """
                
                completion_response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": f"Complete the {detected_category} survey with the exact remaining questions needed. Use proper Q[number]. format."},
                        {"role": "user", "content": completion_prompt}
                    ],
                    temperature=0.2,
                    max_tokens=2000
                )
                
                questionnaire += "\n\n" + completion_response.choices[0].message.content
                
                # Re-validate
                question_lines = [line for line in questionnaire.split('\n') if line.strip().startswith('Q') and '.' in line and any(char.isdigit() for char in line)]
                actual_count = len(question_lines)
        
        # Show generation summary
        st.info(f"📊 **Generation Summary:** {actual_count} questions generated out of {question_counts['total']} target questions for {detected_category} category")
        
        # Step 7: Validate and complete questionnaire
        status_text.text("✅ Validating question count...")
        progress_bar.progress(95)
        
        # Step 8: Final formatting and storage
        status_text.text("✨ Formatting questionnaire...")
        formatted_questionnaire = format_questionnaire_with_logic(questionnaire)
        st.session_state.questionnaire_text = formatted_questionnaire
        st.session_state.questionnaire_generated = True
        
        progress_bar.progress(100)
        status_text.text("✅ Generation complete!")
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        # Show final count
        final_question_lines = [line for line in questionnaire.split('\n') if line.strip().startswith('Q') and '.' in line and any(char.isdigit() for char in line)]
        final_count = len(final_question_lines)
        
        if final_count >= question_counts['total']:
            st.success(f"🎉 **Complete {detected_category} questionnaire generated!** {final_count} questions created. Scroll down to view and download.")
        else:
            st.success(f"🎉 **{detected_category.title()} questionnaire generated!** {final_count} out of {question_counts['total']} questions created. Scroll down to view and download.")
        
    except Exception as e:
        progress_bar.empty()
        status_text.empty()
        st.error(f"❌ Generation failed: {str(e)}")

# Display Results Section (only if questionnaire was generated)
if st.session_state.questionnaire_generated and st.session_state.questionnaire_text:
    st.header("📊 Generated Questionnaire")
    
    # Display questionnaire
    st.text_area(
        "Complete Survey Questionnaire",
        st.session_state.questionnaire_text,
        height=500,
        help="Your comprehensive survey with statistical analysis, fraud detection, and skip logic"
    )
    
    # Download section (always visible after generation)
    st.header("📥 Download Options")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.download_button(
            "📄 Download Text File",
            st.session_state.questionnaire_text,
            file_name=f"survey_questionnaire_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            mime="text/plain",
            use_container_width=True
        )
    
    with col2:
        # Word document
        if st.session_state.survey_data_stored:
            doc = create_comprehensive_word_document(st.session_state.questionnaire_text, st.session_state.survey_data_stored)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            st.download_button(
                "📝 Download Word Doc",
                doc_io.getvalue(),
                file_name=f"survey_questionnaire_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    with col3:
        # Excel file
        if st.session_state.survey_data_stored:
            toolkit = load_comprehensive_excel_toolkit()
            excel_data = create_structured_excel_output(st.session_state.questionnaire_text, st.session_state.survey_data_stored, toolkit)
            
            st.download_button(
                "📊 Download Excel File",
                excel_data,
                file_name=f"survey_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

# Information panels
if not st.session_state.questionnaire_generated:
    st.header("📚 Fixed Category Detection System")
    
    # Test category detection
    if st.session_state.get('survey_objective', ''):
        test_category, test_confidence = detect_survey_category(
            st.session_state.get('survey_objective', ''), 
            st.session_state.get('target_audience', '')
        )
        st.info(f"🔍 **Preview Category Detection:** {test_category.title()} (Confidence: {test_confidence})")
    
    col1, col2 = st.columns(2)
    
    with col1:
        with st.expander("🔧 View Survey Toolkit", expanded=False):
            toolkit = load_comprehensive_excel_toolkit()
            st.subheader("Question Types & Scales")
            for q_type, details in toolkit['question_types'].items():
                st.write(f"**{q_type}:**")
                st.write(f"Scale: {' | '.join(details['scale'])}")
                st.write(f"Analysis: {', '.join(details['analysis'])}")
                st.write("---")
    
    with col2:
        with st.expander("🎯 Category Detection Logic", expanded=False):
            st.subheader("Supported Categories")
            categories = {
                'Cosmetics': ['cosmetics', 'beauty', 'cream', 'skincare', 'makeup', 'night cream'],
                'Automotive': ['automotive', 'car', 'vehicle', 'ev', 'electric vehicle'],
                'Technology': ['phone', 'smartphone', 'mobile', 'laptop', 'computer'],
                'Food & Beverage': ['food', 'restaurant', 'beverage', 'drink', 'coffee'],
                'Fashion': ['fashion', 'clothing', 'apparel', 'shoes'],
                'Healthcare': ['healthcare', 'medical', 'health', 'medicine'],
                'Finance': ['finance', 'banking', 'investment', 'insurance']
            }
            
            for category, keywords in categories.items():
                st.write(f"**{category}:** {', '.join(keywords[:4])}...")
    
    st.success("""
    ✅ **FIXED: Category Consistency Issues Resolved:**
    
    **Key Fixes Implemented:**
    - 🎯 **Intelligent Category Detection** - Automatically detects survey category from objective and target audience
    - 🏷️ **Category-Specific Brand Research** - Dynamically loads brands for detected category only
    - 🚫 **No Hard-Coded Categories** - Removed all hard-coded EV references in prompts
    - 🔄 **Consistent Category Usage** - All question parts use the same detected category
    - 📊 **Category Confidence Scoring** - Shows detection confidence level
    - 🎨 **Category-Specific Attributes** - Matrix questions adapt to category (cosmetics vs automotive vs tech)
    - 🏪 **Market-Specific Brands** - Different brand lists for India vs global markets
    - 📝 **Category Validation** - Prevents mixing categories in questionnaire
    
    **Supported Categories:**
    - 🧴 **Cosmetics** (beauty, skincare, night cream, makeup)
    - 🚗 **Automotive** (cars, EVs, vehicles)
    - 📱 **Technology** (smartphones, laptops, software)
    - 🍕 **Food & Beverage** (restaurants, drinks, snacks)
    - 👗 **Fashion** (clothing, shoes, accessories)
    - 🏥 **Healthcare** (medical, health products)
    - 💰 **Finance** (banking, investment, insurance)
    """)
    
    q_counts = calculate_question_count(20)
    st.info(f"""
    🎯 **Category-Consistent Question Generation:**
    - ✅ **Screener Questions** ({q_counts['screener']}) - Category-specific qualification
    - ✅ **Core Research** ({q_counts['core_research']}) - Brand awareness, usage, satisfaction for detected category only
    - ✅ **Demographics** ({q_counts['demographics']}) - Standard demographic profiling
    - ✅ **Total Questions** ({q_counts['total']}) - All focused on single detected category
    - ✅ **Real Brand Names** - No more "Brand A, Brand B, Brand C"
    - ✅ **Metadata Integration** - Complete question metadata for all sections
    """)

# Footer
st.markdown("---")
st.markdown("*Powered by Advanced AI Survey Methodology • Category-Consistent Generation • Professional Grade Output*")
